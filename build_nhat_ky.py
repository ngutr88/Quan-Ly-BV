from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc=Document(); s=doc.sections[0]
s.top_margin=Inches(.75); s.bottom_margin=Inches(.7); s.left_margin=Inches(.8); s.right_margin=Inches(.8)
styles=doc.styles; n=styles['Normal']; n.font.name='Aptos'; n.font.size=Pt(10.5); n.font.color.rgb=RGBColor(38,50,56); n.paragraph_format.space_after=Pt(5)
for name,size,color in [('Title',24,'0B3954'),('Heading 1',16,'0B3954'),('Heading 2',12,'087E8B')]:
    st=styles[name]; st.font.name='Aptos'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)

def shade(cell,fill):
    p=cell._tc.get_or_add_tcPr(); x=OxmlElement('w:shd'); x.set(qn('w:fill'),fill); p.append(x)
def bullet(t): doc.add_paragraph(t,style='List Bullet')

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(90)
r=p.add_run('NHẬT KÝ TIẾN ĐỘ DỰ ÁN'); r.bold=True; r.font.size=Pt(25); r.font.color.rgb=RGBColor(11,57,84)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('HỆ THỐNG QUẢN LÝ BỆNH VIỆN (HMS)'); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=RGBColor(8,126,139)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Thời gian: 02/06/2026 – 27/07/2026\n08 tuần | Phiên bản 01').font.size=Pt(12)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(80); p.add_run('Tài liệu lập theo PRD của dự án\nNội dung là kế hoạch và tiến độ dự kiến theo tuần').italic=True
doc.add_page_break()

doc.add_heading('1. Thông tin chung',1)
rows=[('Tên dự án','Hệ thống Quản lý Bệnh viện (Hospital Management System - HMS)'),('Mục tiêu','Số hóa quy trình đặt lịch, khám bệnh, kê đơn, quản lý kho, hóa đơn và thanh toán.'),('Đối tượng sử dụng','Admin, Bác sĩ, Bệnh nhân; có khả năng mở rộng cho Lễ tân, Dược sĩ, Kế toán, Điều dưỡng.'),('Phạm vi nhật ký','Phân tích yêu cầu, thiết kế, phát triển luồng cốt lõi, kiểm thử và hoàn thiện.'),('Thời gian','02/06/2026 đến 27/07/2026, chia thành 08 tuần liên tiếp.')]
t=doc.add_table(rows=len(rows),cols=2); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(a,b) in enumerate(rows):
    t.cell(i,0).text=a; t.cell(i,1).text=b; shade(t.cell(i,0),'DCEFF2'); t.cell(i,0).paragraphs[0].runs[0].bold=True

doc.add_heading('2. Nhật ký theo tuần',1)
weeks=[
('Tuần 1','02/06/2026 – 08/06/2026','Khởi động và phân tích yêu cầu',['Đọc PRD; xác định 3 nhóm người dùng và use case ưu tiên UC-01 đến UC-09.','Phân tích quy trình đặt lịch → check-in → khám → kê đơn → trừ kho → hóa đơn → thanh toán.','Xác định yêu cầu RBAC, audit log, bảo mật dữ liệu y tế và soft delete.'],'Chốt phạm vi ban đầu, danh sách module và tiêu chí nghiệp vụ cốt lõi.','Không mở rộng tính năng thấp trước khi luồng cốt lõi ổn định.'),
('Tuần 2','09/06/2026 – 15/06/2026','Thiết kế kiến trúc và dữ liệu',['Phân rã module theo domain: auth, users, doctors, patients, departments, appointments, examinations, prescriptions, medicines, invoices, notifications, reports.','Thiết kế ERD cho người dùng, bác sĩ, bệnh nhân, khoa, lịch khám, phiếu khám, đơn thuốc, thuốc, lô thuốc, hóa đơn và audit log.','Thống nhất trạng thái lịch khám, hóa đơn và timezone Asia/Ho_Chi_Minh.'],'Có mô hình dữ liệu nền tảng và ranh giới module rõ ràng.','Số slot, thời lượng khám và ngưỡng tồn kho phải cấu hình được.'),
('Tuần 3','16/06/2026 – 22/06/2026','Xác thực, phân quyền và tài khoản',['Thiết kế đăng nhập, đăng ký bệnh nhân và OTP có thời hạn, giới hạn nhập sai, cooldown gửi lại.','Xây dựng RBAC cho Admin, Bác sĩ, Bệnh nhân; kiểm tra quyền ở server.','Bổ sung khóa/mở khóa tài khoản và audit log thao tác nhạy cảm.'],'Hoàn thiện nền tảng auth/RBAC và nguyên tắc lọc dữ liệu theo quyền sở hữu hoặc phân công.','Không log mật khẩu, token, OTP hoặc dữ liệu y tế không cần thiết.'),
('Tuần 4','23/06/2026 – 29/06/2026','Đặt lịch và quản lý lịch khám',['Phát triển tìm kiếm khoa/bác sĩ, chọn giờ và tạo lịch khám.','Kiểm tra trùng lịch, lịch làm việc, ngày nghỉ, số bệnh nhân tối đa và thời lượng khám.','Bổ sung xác nhận, dời, hủy lịch; lưu lý do và chuẩn bị thông báo.'],'Hoàn thiện luồng UC-01 và UC-08 ở mức chức năng chính.','Tình huống đồng thời tạo lịch cần ràng buộc dữ liệu và transaction.'),
('Tuần 5','30/06/2026 – 06/07/2026','Khám bệnh và hồ sơ bệnh nhân',['Liên kết phiếu khám với lịch; nhập triệu chứng, sinh hiệu, chẩn đoán, cận lâm sàng, kết quả và lời dặn.','Tổng hợp tiền sử, dị ứng, lịch sử khám, đơn thuốc và hóa đơn trong hồ sơ bệnh nhân.','Giới hạn bác sĩ theo bệnh nhân/lịch được phân công; bệnh nhân không sửa dữ liệu chuyên môn.'],'Hoàn thiện luồng khám bệnh và hồ sơ liên tục, đáp ứng UC-02 và UC-09.','Kiểm tra kỹ phạm vi truy cập và nhật ký truy vết dữ liệu y tế.'),
('Tuần 6','07/07/2026 – 13/07/2026','Kê đơn và quản lý kho thuốc',['Kê đơn với liều dùng, cách dùng, số lượng và số ngày dùng.','Kiểm tra tồn kho, dị ứng và cảnh báo tương tác.','Trừ kho theo FEFO; không tồn âm nếu chưa có override có quyền và lý do; cảnh báo dưới ngưỡng/sắp hết hạn.'],'Kết nối được đơn thuốc với kho, tạo nền tảng cho UC-03 và UC-05.','Kiểm thử nhiều lô, lô hết hạn và trường hợp thiếu số lượng.'),
('Tuần 7','14/07/2026 – 20/07/2026','Hóa đơn, thanh toán và thông báo',['Sinh hóa đơn từ phiên khám, thuốc và dịch vụ; quản lý trạng thái thanh toán.','Thiết kế thanh toán tại quầy và online; callback xác thực, idempotent; lỗi/timeout giữ chưa thanh toán.','Chuẩn bị thông báo lịch, nhắc lịch, thanh toán và cảnh báo kho.'],'Hoàn thiện UC-04 và nền tảng thông báo; thay đổi hóa đơn có audit log.','Mô phỏng callback lặp, sai chữ ký và thanh toán thất bại.'),
('Tuần 8','21/07/2026 – 27/07/2026','Tích hợp, kiểm thử và tổng kết',['Chạy kiểm thử end-to-end từ đăng ký, đặt lịch, khám, kê đơn, kho, hóa đơn đến thanh toán.','Kiểm thử phân quyền, xung đột lịch, FEFO, validation form, callback và quyền tải file.','Rà soát loading/empty/error/permission denied, responsive và tài liệu bàn giao.'],'Có bản nghiệm thu nội bộ theo use case ưu tiên; lập danh sách lỗi và lộ trình sau MVP.','BHYT, app mobile, telemedicine và phân tích nâng cao để giai đoạn sau.')]
for title,date,focus,tasks,result,risk in weeks:
    doc.add_heading(title+' | '+date,2); p=doc.add_paragraph(); p.add_run('Trọng tâm: ').bold=True; p.add_run(focus)
    p=doc.add_paragraph(); p.add_run('Công việc thực hiện').bold=True
    for x in tasks: bullet(x)
    p=doc.add_paragraph(); p.add_run('Kết quả / sản phẩm: ').bold=True; p.add_run(result)
    p=doc.add_paragraph(); p.add_run('Khó khăn / lưu ý: ').bold=True; p.add_run(risk)

doc.add_heading('3. Tổng kết sau 8 tuần',1)
for x in ['Đã bao phủ các ưu tiên cao của PRD: đặt lịch, khám, kê đơn, hóa đơn–thanh toán, kho thuốc, đăng ký, đăng nhập–phân quyền và hồ sơ bệnh nhân.','Các luồng quan trọng cần transaction, kiểm tra quyền ở server và audit log để bảo đảm an toàn dữ liệu.','Giai đoạn tiếp theo: hoàn thiện kiểm thử tích hợp, triển khai, báo cáo nâng cao, cổng thanh toán thật và mở rộng vai trò vận hành.']: bullet(x)
doc.add_heading('4. Ghi chú',1); doc.add_paragraph('Nhật ký được lập theo kế hoạch 8 tuần dựa trên PRD. Nếu nhóm có biên bản họp, commit hoặc task thực tế theo ngày, có thể cập nhật các mục công việc và kết quả để phản ánh chính xác lịch sử triển khai.')
f=s.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; f.add_run('HMS | Nhật ký tiến độ 8 tuần | 02/06/2026 – 27/07/2026').font.size=Pt(9)
doc.core_properties.title='Nhật ký tiến độ dự án HMS - 8 tuần'; doc.core_properties.author=''; doc.save('Nhat_ky_tien_do_HMS_8_tuan.docx')
